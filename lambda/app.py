import base64
import hashlib
import html.parser
import ipaddress
import json
import logging
import os
import re
import urllib.error
import urllib.parse
import urllib.request
from collections.abc import Sequence
from io import BytesIO
from typing import Any

import boto3


logger = logging.getLogger(__name__)
s3_client = boto3.client("s3")

RESUME_BUCKET_ENV = "RESUME_BUCKET"
RESUME_KEY_ENV = "RESUME_KEY"
SEMANTIC_MATCHING_ENABLED_ENV = "SEMANTIC_MATCHING_ENABLED"
SEMANTIC_EMBEDDING_PROVIDER_ENV = "SEMANTIC_EMBEDDING_PROVIDER"
SEMANTIC_MODEL_NAME_ENV = "SEMANTIC_MODEL_NAME"
BEDROCK_EMBEDDING_MODEL_ID_ENV = "BEDROCK_EMBEDDING_MODEL_ID"
BEDROCK_EMBEDDING_DIMENSIONS_ENV = "BEDROCK_EMBEDDING_DIMENSIONS"
EMBEDDING_CACHE_BUCKET_ENV = "EMBEDDING_CACHE_BUCKET"
EMBEDDING_CACHE_PREFIX_ENV = "EMBEDDING_CACHE_PREFIX"

DEFAULT_SEMANTIC_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"
DEFAULT_SEMANTIC_EMBEDDING_PROVIDER = "bedrock"
LOCAL_EMBEDDING_PROVIDER = "local"
BEDROCK_EMBEDDING_PROVIDER = "bedrock"
DEFAULT_BEDROCK_EMBEDDING_MODEL_ID = "amazon.titan-embed-text-v2:0"
DEFAULT_BEDROCK_EMBEDDING_DIMENSIONS = 512
DEFAULT_EMBEDDING_CACHE_PREFIX = "embeddings/resume"
EMBEDDING_CACHE_SCHEMA_VERSION = "1.0"
DEFAULT_KEYWORD_SCORE_WEIGHT = 0.45
DEFAULT_SEMANTIC_SCORE_WEIGHT = 0.55
MATCHED_REQUIREMENT_MIN_SCORE = 40
TOP_EVIDENCE_DIAGNOSTIC_LIMIT = 3
MAX_RESUME_EVIDENCE_WINDOW_WORDS = 80
MIN_RESUME_EVIDENCE_WINDOW_WORDS = 8
LONG_RESUME_CHUNK_WORDS = 45
MAX_JOB_DESCRIPTION_URL_BYTES = 1_000_000
JOB_DESCRIPTION_URL_TIMEOUT_SECONDS = 5
SUPPORTED_RESUME_EXTENSIONS = {".txt", ".pdf", ".docx"}
SUPPORTED_JOB_FILE_EXTENSIONS = {".txt", ".md"}
KEYWORD_TOKEN_PATTERN = re.compile(r"[a-z0-9][a-z0-9+#./-]*")
KEYWORD_TRAILING_PUNCTUATION = ".,;:!?)]}"
CONTRACTION_FRAGMENTS = {"d", "ll", "m", "re", "s", "t", "ve"}
LOW_VALUE_KEYWORDS = {"responsibilities", "responsibility", "role", "responses", "systems"}
KEYWORD_ALIASES = {"apis": "api"}
WEAK_REQUIREMENT_HEADINGS = {
    "about",
    "about us",
    "benefits",
    "description",
    "job description",
    "overview",
    "preferred qualifications",
    "qualifications",
    "requirements",
    "responsibilities",
    "role",
    "summary",
    "what you will do",
}
LOW_VALUE_REQUIREMENT_PATTERNS = [
    r"\b(?:job|requisition|req|posting)\s*(?:id|number|no\.?)\b",
    r"\b(?:base\s+pay|salary|compensation|pay\s+range|hourly\s+range)\b",
    r"\b(?:benefits?|health\s+insurance|dental|vision|401k|paid\s+time\s+off)\b",
    r"\b(?:equal\s+opportunity|eeo|reasonable\s+accommodation)\b",
    r"\b(?:applicants?\s+with\s+arrest|criminal\s+histories|background\s+check)\b",
    r"\b(?:location|work\s+location|workplace\s+type)\s*:",
    r"\b(?:remote|hybrid|onsite)\s+(?:role|position|work|schedule)\b",
    r"\b(?:about\s+us|about\s+the\s+company|our\s+company|who\s+we\s+are)\b",
]
EVIDENCE_ALIAS_PATTERNS = {
    "program_management": [
        r"\bprogram\s*/\s*project\s+management\b",
        r"\bproject\s*/\s*program\s+management\b",
        r"\bprogram\s+management\b",
    ],
    "project_management": [
        r"\bprogram\s*/\s*project\s+management\b",
        r"\bproject\s*/\s*program\s+management\b",
        r"\bproject\s+management\b",
    ],
    "ai_ml": [
        r"\bai\s*/\s*ml\b",
        r"\bml\s*/\s*ai\b",
        r"\bartificial\s+intelligence\b",
        r"\bmachine\s+learning\b",
    ],
    "agentic_ai": [
        r"\bagentic\s+ai\b",
        r"\bai\s+agent(?:s)?\b",
        r"\bagent(?:ic)?\s+tool(?:s)?\b",
    ],
    "content_data_systems": [
        r"\bcontent\s*/\s*data\s+systems\b",
        r"\bcontent\s+systems\b",
        r"\bdata\s+systems\b",
    ],
}

_embedding_provider = None
_local_embedding_model = None
bedrock_runtime_client = None

STOP_WORDS = {
    "a",
    "about",
    "above",
    "after",
    "again",
    "against",
    "all",
    "am",
    "an",
    "and",
    "any",
    "are",
    "as",
    "at",
    "be",
    "because",
    "been",
    "before",
    "being",
    "below",
    "between",
    "both",
    "but",
    "by",
    "can",
    "did",
    "do",
    "does",
    "doing",
    "down",
    "during",
    "each",
    "few",
    "for",
    "from",
    "further",
    "had",
    "has",
    "have",
    "having",
    "he",
    "her",
    "here",
    "hers",
    "herself",
    "him",
    "himself",
    "his",
    "how",
    "i",
    "if",
    "in",
    "into",
    "is",
    "it",
    "its",
    "itself",
    "just",
    "me",
    "more",
    "most",
    "my",
    "myself",
    "no",
    "nor",
    "not",
    "now",
    "of",
    "off",
    "on",
    "once",
    "only",
    "or",
    "other",
    "our",
    "ours",
    "ourselves",
    "out",
    "over",
    "own",
    "same",
    "she",
    "should",
    "so",
    "some",
    "such",
    "than",
    "that",
    "the",
    "their",
    "theirs",
    "them",
    "themselves",
    "then",
    "there",
    "these",
    "they",
    "this",
    "those",
    "through",
    "to",
    "too",
    "under",
    "until",
    "up",
    "very",
    "was",
    "we",
    "were",
    "what",
    "when",
    "where",
    "which",
    "while",
    "who",
    "whom",
    "why",
    "will",
    "with",
    "you",
    "your",
    "yours",
    "yourself",
    "yourselves",
}


def lambda_handler(event: dict[str, Any], context: Any) -> dict[str, Any]:
    request_method = _request_method(event)
    if request_method and request_method != "POST":
        return _response(405, {"message": "Method not allowed"})

    try:
        payload = _parse_body(event)
        job_description = _job_description_from_payload(payload)
        resume_object = _resume_object_from_payload(payload)
    except (KeyError, TypeError, ValueError, json.JSONDecodeError) as exc:
        return _response(400, {"message": f"Invalid request body: {exc}"})
    except urllib.error.URLError:
        logger.exception("Failed to read job description URL")
        return _response(502, {"message": "Unable to read job description URL"})

    if not isinstance(job_description, str) or not job_description.strip():
        return _response(400, {"message": "job_description must be a non-empty string"})

    if resume_object is None:
        try:
            resume_object = _read_resume_object_from_s3()
        except ValueError as exc:
            return _response(500, {"message": str(exc)})
        except Exception:
            logger.exception("Failed to read resume from S3")
            return _response(502, {"message": "Unable to read resume from S3"})

    result = compare_resume_to_job(
        resume_object["text"],
        job_description,
        resume_source=resume_object.get("source"),
    )
    return _response(200, result)


class EmbeddingProvider:
    model_id: str
    dimensions: int | None
    normalize: bool

    def encode(self, text: str) -> list[float]:
        raise NotImplementedError


class LocalSentenceTransformerEmbeddingProvider(EmbeddingProvider):
    def __init__(self, model: Any, model_id: str):
        self.model = model
        self.model_id = model_id
        self.dimensions = None
        self.normalize = False

    def encode(self, text: str) -> list[float]:
        return list(self.model.encode(text))


class BedrockEmbeddingProvider(EmbeddingProvider):
    def __init__(
        self,
        client: Any,
        model_id: str,
        dimensions: int,
        normalize: bool = True,
    ):
        self.client = client
        self.model_id = model_id
        self.dimensions = dimensions
        self.normalize = normalize

    def encode(self, text: str) -> list[float]:
        response = self.client.invoke_model(
            modelId=self.model_id,
            body=json.dumps(
                {
                    "inputText": text,
                    "dimensions": self.dimensions,
                    "normalize": self.normalize,
                }
            ),
            accept="application/json",
            contentType="application/json",
        )
        payload = json.loads(response["body"].read().decode("utf-8"))
        return payload["embedding"]


class S3EmbeddingCache:
    def __init__(self, client: Any, bucket: str, prefix: str):
        self.client = client
        self.bucket = bucket
        self.prefix = prefix.strip("/")

    def cache_key(
        self,
        resume_source: dict[str, str],
        provider: EmbeddingProvider,
    ) -> str:
        cache_identity = {
            "schema_version": EMBEDDING_CACHE_SCHEMA_VERSION,
            "resume_bucket": resume_source["bucket"],
            "resume_key": resume_source["key"],
            "resume_etag": resume_source["etag"],
            "model_id": provider.model_id,
            "dimensions": provider.dimensions,
            "normalize": provider.normalize,
        }
        digest = hashlib.sha256(
            json.dumps(cache_identity, sort_keys=True).encode("utf-8")
        ).hexdigest()
        return f"{self.prefix}/{digest}.json"

    def get(self, key: str) -> list[float] | None:
        try:
            s3_object = self.client.get_object(Bucket=self.bucket, Key=key)
        except Exception as exc:
            if _is_s3_cache_miss(exc):
                return None
            raise

        payload = json.loads(s3_object["Body"].read().decode("utf-8"))
        return payload["embedding"]["vector"]

    def put(
        self,
        key: str,
        resume_source: dict[str, str],
        provider: EmbeddingProvider,
        embedding: Sequence[float],
    ) -> None:
        payload = {
            "schema_version": EMBEDDING_CACHE_SCHEMA_VERSION,
            "source": resume_source,
            "embedding": {
                "provider": _semantic_embedding_provider_name(),
                "model_id": provider.model_id,
                "dimensions": provider.dimensions,
                "normalize": provider.normalize,
                "vector": list(embedding),
            },
        }
        self.client.put_object(
            Bucket=self.bucket,
            Key=key,
            Body=json.dumps(payload).encode("utf-8"),
            ContentType="application/json",
        )


def compare_resume_to_job(
    resume_text: str,
    job_description: str,
    resume_source: dict[str, str] | None = None,
) -> dict[str, Any]:
    resume_keywords = _extract_keywords(resume_text)
    job_keywords = _extract_keywords(job_description)

    if not job_keywords:
        return {"score": 0, "matching_keywords": [], "missing_keywords": []}

    matching_keywords = sorted(job_keywords & resume_keywords)
    missing_keywords = sorted(job_keywords - resume_keywords)
    keyword_score = calculate_keyword_score(resume_keywords, job_keywords)

    if not _semantic_matching_enabled():
        return {
            "score": keyword_score,
            "matching_keywords": matching_keywords,
            "missing_keywords": missing_keywords,
        }

    embedding_provider = _load_embedding_provider()
    embedding_cache = _load_embedding_cache() if resume_source else None
    semantic_score = calculate_semantic_score(
        resume_text,
        job_description,
        embedding_model=embedding_provider,
        resume_source=resume_source,
        embedding_cache=embedding_cache,
    )
    chunked_semantic_score = calculate_chunked_semantic_score(
        resume_text,
        job_description,
        embedding_model=embedding_provider,
    )
    fit_analysis = build_fit_analysis(
        resume_text,
        job_description,
        embedding_model=embedding_provider,
    )
    score = combine_scores(keyword_score, semantic_score)

    return {
        "score": score,
        "keyword_score": keyword_score,
        "semantic_score": semantic_score,
        "chunked_semantic_score": chunked_semantic_score,
        "matching_keywords": matching_keywords,
        "missing_keywords": missing_keywords,
        "matched_requirements": fit_analysis["matched_requirements"],
        "gaps": fit_analysis["gaps"],
        "semantic_model": embedding_provider.model_id,
        "semantic_provider": _semantic_embedding_provider_name(),
        "weights": {
            "keyword": DEFAULT_KEYWORD_SCORE_WEIGHT,
            "semantic": DEFAULT_SEMANTIC_SCORE_WEIGHT,
        },
    }


def calculate_keyword_score(resume_keywords: set[str], job_keywords: set[str]) -> int:
    if not job_keywords:
        return 0

    matching_keywords = job_keywords & resume_keywords
    return round((len(matching_keywords) / len(job_keywords)) * 100)


def calculate_semantic_score(
    resume_text: str,
    job_description: str,
    embedding_model: Any | None = None,
    resume_source: dict[str, str] | None = None,
    embedding_cache: S3EmbeddingCache | None = None,
) -> int:
    model = embedding_model or _load_embedding_provider()
    resume_embedding = _get_resume_embedding(
        resume_text,
        resume_source,
        model,
        embedding_cache,
    )
    job_embedding = model.encode(job_description)
    similarity = max(0.0, cosine_similarity(resume_embedding, job_embedding))
    return round(similarity * 100)


def calculate_chunked_semantic_score(
    resume_text: str,
    job_description: str,
    embedding_model: Any | None = None,
) -> int:
    model = embedding_model or _load_embedding_provider()
    resume_chunks = _chunk_resume_text(resume_text)
    job_chunks = _chunk_job_description_text(job_description)

    if not resume_chunks or not job_chunks:
        return 0

    resume_embeddings = [model.encode(chunk) for chunk in resume_chunks]
    job_embeddings = [model.encode(chunk) for chunk in job_chunks]
    best_matches = [
        max(
            max(0.0, cosine_similarity(job_embedding, resume_embedding))
            for resume_embedding in resume_embeddings
        )
        for job_embedding in job_embeddings
    ]

    return round((sum(best_matches) / len(best_matches)) * 100)


def build_fit_analysis(
    resume_text: str,
    job_description: str,
    embedding_model: Any | None = None,
) -> dict[str, Any]:
    requirement_matches = _score_requirement_evidence(
        _extract_requirement_candidates(job_description),
        _extract_resume_evidence_candidates(resume_text),
        embedding_model,
    )

    matched_requirements = []
    gaps = []
    for match in requirement_matches:
        if match["score"] >= MATCHED_REQUIREMENT_MIN_SCORE and match["evidence"]:
            matched_requirements.append(
                {
                    "requirement": match["requirement"],
                    "score": match["score"],
                    "evidence": match["evidence"],
                }
            )
        else:
            gaps.append(
                {
                    "requirement": match["requirement"],
                    "score": match["score"],
                }
            )

    score_summary = _fit_analysis_score_summary(
        requirement_matches,
        matched_requirements,
        gaps,
    )
    logger.info("Fit analysis score summary: %s", json.dumps(score_summary))

    return {
        "matched_requirements": matched_requirements,
        "gaps": gaps,
        "_requirement_evidence_scores": requirement_matches,
        "_score_summary": score_summary,
    }


def _fit_analysis_score_summary(
    requirement_matches: Sequence[dict[str, Any]],
    matched_requirements: Sequence[dict[str, Any]],
    gaps: Sequence[dict[str, Any]],
) -> dict[str, Any]:
    scores = [match["score"] for match in requirement_matches]
    if not scores:
        return {
            "requirement_count": 0,
            "matched_count": 0,
            "gap_count": 0,
            "matched_requirement_min_score": MATCHED_REQUIREMENT_MIN_SCORE,
            "min_score": None,
            "max_score": None,
            "average_score": None,
        }

    return {
        "requirement_count": len(requirement_matches),
        "matched_count": len(matched_requirements),
        "gap_count": len(gaps),
        "matched_requirement_min_score": MATCHED_REQUIREMENT_MIN_SCORE,
        "min_score": min(scores),
        "max_score": max(scores),
        "average_score": round(sum(scores) / len(scores)),
    }


def _score_requirement_evidence(
    requirements: Sequence[str],
    evidence_chunks: Sequence[Any],
    embedding_model: Any | None = None,
) -> list[dict[str, Any]]:
    if not requirements:
        return []

    evidence_candidates = []
    for evidence in evidence_chunks:
        candidate = _coerce_evidence_candidate(evidence)
        if candidate["evidence"]:
            evidence_candidates.append(candidate)
    evidence_texts = [
        candidate["evidence"]
        for candidate in evidence_candidates
    ]
    evidence_keywords = {
        evidence: _extract_keywords(evidence)
        for evidence in evidence_texts
    }
    evidence_alias_concepts = {
        evidence: _extract_evidence_alias_concepts(evidence)
        for evidence in evidence_texts
    }
    evidence_embeddings = {}
    if embedding_model is not None:
        evidence_embeddings = {
            evidence: embedding_model.encode(evidence)
            for evidence in evidence_texts
        }

    matches = []
    for requirement in requirements:
        requirement_keywords = _extract_keywords(requirement)
        requirement_embedding = (
            embedding_model.encode(requirement)
            if embedding_model is not None and evidence_candidates
            else None
        )
        best_match = {
            "requirement": requirement,
            "score": 0,
            "evidence": "",
            "keyword_score": 0,
            "semantic_score": None,
            "alias_score": 0,
            "matching_keywords": [],
            "matching_aliases": [],
            "top_evidence": [],
        }
        requirement_alias_concepts = _extract_evidence_alias_concepts(requirement)
        evidence_rankings = []

        for candidate in evidence_candidates:
            evidence = candidate["evidence"]
            matching_keywords = requirement_keywords & evidence_keywords[evidence]
            keyword_score = (
                round((len(matching_keywords) / len(requirement_keywords)) * 100)
                if requirement_keywords
                else 0
            )
            matching_aliases = (
                requirement_alias_concepts & evidence_alias_concepts[evidence]
            )
            alias_score = (
                round((len(matching_aliases) / len(requirement_alias_concepts)) * 100)
                if requirement_alias_concepts
                else 0
            )
            semantic_score = None
            score = keyword_score
            if requirement_embedding is not None:
                semantic_score = round(
                    max(
                        0.0,
                        cosine_similarity(
                            requirement_embedding,
                            evidence_embeddings[evidence],
                        ),
                    )
                    * 100
                )
                score = _rank_requirement_evidence_score(
                    keyword_score,
                    semantic_score,
                    alias_score,
                )
            elif alias_score:
                score = _rank_requirement_evidence_score(
                    keyword_score,
                    None,
                    alias_score,
                )

            evidence_ranking = {
                "evidence": evidence,
                "score": score,
                "keyword_score": keyword_score,
                "semantic_score": semantic_score,
                "alias_score": alias_score,
                "matching_keywords": sorted(matching_keywords),
                "matching_aliases": sorted(matching_aliases),
                "evidence_type": candidate["type"],
                "source_chunks": candidate["source_chunks"],
                "word_count": _word_count(evidence),
            }
            evidence_rankings.append(evidence_ranking)

        ranked_evidence = sorted(
            evidence_rankings,
            key=_evidence_rerank_key,
            reverse=True,
        )
        if ranked_evidence:
            selected_evidence = ranked_evidence[0]
            best_match = {
                "requirement": requirement,
                "score": selected_evidence["score"],
                "evidence": selected_evidence["evidence"],
                "keyword_score": selected_evidence["keyword_score"],
                "semantic_score": selected_evidence["semantic_score"],
                "alias_score": selected_evidence["alias_score"],
                "matching_keywords": selected_evidence["matching_keywords"],
                "matching_aliases": selected_evidence["matching_aliases"],
                "evidence_type": selected_evidence["evidence_type"],
                "source_chunks": selected_evidence["source_chunks"],
                "top_evidence": [],
            }

        best_match["top_evidence"] = ranked_evidence[:TOP_EVIDENCE_DIAGNOSTIC_LIMIT]
        matches.append(best_match)

    return matches


def _coerce_evidence_candidate(evidence: Any) -> dict[str, Any]:
    if isinstance(evidence, dict):
        evidence_text = evidence.get("evidence", "")
        return {
            "evidence": evidence_text,
            "type": evidence.get("type", "chunk"),
            "source_chunks": evidence.get("source_chunks", [evidence_text]),
        }

    return {
        "evidence": str(evidence),
        "type": "chunk",
        "source_chunks": [str(evidence)],
    }


def _evidence_rerank_key(ranking: dict[str, Any]) -> tuple[int, int, int, int, int, int]:
    return (
        ranking["score"],
        ranking["alias_score"],
        ranking["keyword_score"],
        _evidence_length_quality_score(ranking["word_count"]),
        1 if ranking["evidence_type"] == "window" else 0,
        -ranking["word_count"],
    )


def _evidence_length_quality_score(word_count: int) -> int:
    if 8 <= word_count <= 60:
        return 3
    if 5 <= word_count <= MAX_RESUME_EVIDENCE_WINDOW_WORDS:
        return 2
    if word_count <= 100:
        return 1
    return 0


def _rank_requirement_evidence_score(
    keyword_score: int,
    semantic_score: int | None,
    alias_score: int,
) -> int:
    if semantic_score is None:
        return min(100, keyword_score + round(alias_score * 0.20))

    baseline_score = combine_scores(keyword_score, semantic_score)
    return min(100, baseline_score + round(alias_score * 0.20))


def _extract_evidence_alias_concepts(text: str) -> set[str]:
    normalized_text = _normalize_text(text)
    return {
        concept
        for concept, patterns in EVIDENCE_ALIAS_PATTERNS.items()
        if any(re.search(pattern, normalized_text) for pattern in patterns)
    }


def _extract_requirement_candidates(job_description: str) -> list[str]:
    candidates = []
    seen = set()
    for chunk in _chunk_job_description_text(job_description):
        candidate = _normalize_requirement_candidate(chunk)
        if not candidate or candidate.casefold() in seen:
            continue
        seen.add(candidate.casefold())
        candidates.append(candidate)

    return candidates


def _normalize_requirement_candidate(text: str) -> str:
    candidate = _normalize_chunk_text(text).strip(" -:;")
    if not candidate:
        return ""

    normalized = candidate.casefold()
    if normalized in WEAK_REQUIREMENT_HEADINGS:
        return ""
    if _is_low_value_requirement_candidate(candidate):
        return ""

    keywords = _extract_keywords(candidate)
    if len(candidate) < 8 or len(keywords) < 2:
        return ""

    return candidate


def _is_low_value_requirement_candidate(candidate: str) -> bool:
    normalized = _normalize_text(candidate)
    if any(re.search(pattern, normalized) for pattern in LOW_VALUE_REQUIREMENT_PATTERNS):
        return True

    keywords = _extract_keywords(candidate)
    if len(candidate) > 180 and len(keywords) < 8:
        return True

    return False


def _extract_resume_evidence_chunks(resume_text: str) -> list[str]:
    return _chunk_resume_text(resume_text)


def _extract_resume_evidence_candidates(resume_text: str) -> list[dict[str, Any]]:
    base_chunks = _chunk_resume_text(resume_text)
    candidates = []
    seen = set()

    def add_candidate(evidence: str, candidate_type: str, source_chunks: list[str]) -> None:
        normalized = _normalize_chunk_text(evidence)
        if not normalized or normalized.casefold() in seen:
            return
        seen.add(normalized.casefold())
        candidates.append(
            {
                "evidence": normalized,
                "type": candidate_type,
                "source_chunks": source_chunks,
            }
        )

    for chunk in base_chunks:
        add_candidate(chunk, "chunk", [chunk])
        if _word_count(chunk) >= LONG_RESUME_CHUNK_WORDS:
            sentence_chunks = _split_sentence_chunks(chunk)
            for sentence in sentence_chunks:
                add_candidate(sentence, "sentence", [chunk])
            for first, second in zip(sentence_chunks, sentence_chunks[1:]):
                window = f"{first} {second}"
                if _is_reasonable_evidence_window(window):
                    add_candidate(window, "window", [first, second])

    for first, second in zip(base_chunks, base_chunks[1:]):
        window = f"{first} {second}"
        if _is_reasonable_evidence_window(window) and _should_window_adjacent_chunks(
            first,
            second,
        ):
            add_candidate(window, "window", [first, second])

    return candidates


def _split_sentence_chunks(text: str) -> list[str]:
    return [
        normalized
        for sentence in re.split(r"(?<=[.!?])\s+", text)
        if (normalized := _normalize_chunk_text(sentence))
    ]


def _is_reasonable_evidence_window(text: str) -> bool:
    word_count = _word_count(text)
    return MIN_RESUME_EVIDENCE_WINDOW_WORDS <= word_count <= MAX_RESUME_EVIDENCE_WINDOW_WORDS


def _should_window_adjacent_chunks(first: str, second: str) -> bool:
    first_keywords = _extract_keywords(first)
    second_keywords = _extract_keywords(second)
    if first_keywords & second_keywords:
        return True

    first_aliases = _extract_evidence_alias_concepts(first)
    second_aliases = _extract_evidence_alias_concepts(second)
    if first_aliases & second_aliases:
        return True

    return first.rstrip().endswith(":") and _word_count(first) <= 8


def _word_count(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text))


def _chunk_resume_text(text: str) -> list[str]:
    chunks = []
    for block in _split_text_blocks(text):
        bullet_chunks = _split_bullet_chunks(block)
        if bullet_chunks:
            chunks.extend(bullet_chunks)
        elif normalized := _normalize_chunk_text(block):
            chunks.append(normalized)

    return chunks



def _chunk_job_description_text(text: str) -> list[str]:
    chunks = []
    for block in _split_text_blocks(text):
        bullet_chunks = _split_bullet_chunks(block)
        if bullet_chunks:
            chunks.extend(bullet_chunks)
        elif normalized := _normalize_chunk_text(block):
            chunks.append(normalized)

    return chunks


def _split_paragraph_chunks(text: str) -> list[str]:
    return [
        normalized
        for part in _split_text_blocks(text)
        if (normalized := _normalize_chunk_text(part))
    ]


def _split_text_blocks(text: str) -> list[str]:
    return re.split(r"\n\s*\n+", text)


def _normalize_chunk_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _split_bullet_chunks(text: str) -> list[str]:
    bullet_matches = list(
        re.finditer(
            r"(?ms)^\s*(?:[-*+]|\d+[.)])\s+(.+?)(?=^\s*(?:[-*+]|\d+[.)])\s+|\Z)",
            text,
        )
    )
    return [
        normalized
        for match in bullet_matches
        if (normalized := _normalize_chunk_text(match.group(1)))
    ]


def _get_resume_embedding(
    resume_text: str,
    resume_source: dict[str, str] | None,
    provider: EmbeddingProvider,
    embedding_cache: S3EmbeddingCache | None,
) -> list[float]:
    if resume_source is None or embedding_cache is None:
        return provider.encode(resume_text)

    cache_key = embedding_cache.cache_key(resume_source, provider)
    cached_embedding = embedding_cache.get(cache_key)
    if cached_embedding is not None:
        return cached_embedding

    resume_embedding = provider.encode(resume_text)
    embedding_cache.put(cache_key, resume_source, provider, resume_embedding)
    return resume_embedding


def combine_scores(
    keyword_score: int | float,
    semantic_score: int | float,
    keyword_weight: float = DEFAULT_KEYWORD_SCORE_WEIGHT,
    semantic_weight: float = DEFAULT_SEMANTIC_SCORE_WEIGHT,
) -> int:
    total_weight = keyword_weight + semantic_weight
    if total_weight <= 0:
        raise ValueError("score weights must sum to a positive value")

    weighted_score = (
        (keyword_score * keyword_weight) + (semantic_score * semantic_weight)
    ) / total_weight
    return round(weighted_score)


def cosine_similarity(
    left_embedding: Sequence[float],
    right_embedding: Sequence[float],
) -> float:
    if len(left_embedding) != len(right_embedding):
        raise ValueError("embeddings must have the same dimension")

    dot_product = sum(left * right for left, right in zip(left_embedding, right_embedding))
    left_magnitude = sum(value * value for value in left_embedding) ** 0.5
    right_magnitude = sum(value * value for value in right_embedding) ** 0.5

    if left_magnitude == 0 or right_magnitude == 0:
        return 0.0

    return dot_product / (left_magnitude * right_magnitude)


def _parse_body(event: dict[str, Any]) -> dict[str, Any]:
    body = event.get("body")
    if body is None:
        raise ValueError("body is required")

    if event.get("isBase64Encoded"):
        body = base64.b64decode(body).decode("utf-8")

    if isinstance(body, str):
        return json.loads(body)

    if isinstance(body, dict):
        return body

    raise ValueError("body must be JSON")


def _request_method(event: dict[str, Any]) -> str | None:
    if event.get("httpMethod"):
        return event["httpMethod"].upper()

    return (
        event.get("requestContext", {})
        .get("http", {})
        .get("method", "")
        .upper()
        or None
    )


def _read_resume_from_s3() -> str:
    return _read_resume_object_from_s3()["text"]


def _read_resume_object_from_s3() -> dict[str, Any]:
    bucket = os.environ.get(RESUME_BUCKET_ENV)
    key = os.environ.get(RESUME_KEY_ENV)

    if not bucket or not key:
        raise ValueError(f"{RESUME_BUCKET_ENV} and {RESUME_KEY_ENV} must be configured")

    s3_object = s3_client.get_object(Bucket=bucket, Key=key)
    etag = s3_object.get("ETag", "").strip('"') or "unknown"
    body = s3_object["Body"].read()
    return {
        "text": _extract_text_from_supported_document(
            body,
            key,
            SUPPORTED_RESUME_EXTENSIONS,
            "resume",
        ),
        "source": {
            "bucket": bucket,
            "key": key,
            "etag": etag,
        },
    }


def _semantic_matching_enabled() -> bool:
    return os.environ.get(SEMANTIC_MATCHING_ENABLED_ENV, "false").casefold() in {
        "1",
        "true",
        "yes",
        "on",
    }


def _semantic_model_name() -> str:
    return os.environ.get(SEMANTIC_MODEL_NAME_ENV, DEFAULT_SEMANTIC_MODEL_NAME)


def _semantic_embedding_provider_name() -> str:
    return os.environ.get(
        SEMANTIC_EMBEDDING_PROVIDER_ENV,
        DEFAULT_SEMANTIC_EMBEDDING_PROVIDER,
    ).casefold()


def _bedrock_embedding_model_id() -> str:
    return os.environ.get(
        BEDROCK_EMBEDDING_MODEL_ID_ENV,
        DEFAULT_BEDROCK_EMBEDDING_MODEL_ID,
    )


def _bedrock_embedding_dimensions() -> int:
    value = os.environ.get(
        BEDROCK_EMBEDDING_DIMENSIONS_ENV,
        str(DEFAULT_BEDROCK_EMBEDDING_DIMENSIONS),
    )
    try:
        dimensions = int(value)
    except ValueError as exc:
        raise ValueError(f"{BEDROCK_EMBEDDING_DIMENSIONS_ENV} must be an integer") from exc

    if dimensions <= 0:
        raise ValueError(f"{BEDROCK_EMBEDDING_DIMENSIONS_ENV} must be positive")

    return dimensions


def _embedding_cache_bucket() -> str | None:
    return os.environ.get(EMBEDDING_CACHE_BUCKET_ENV) or os.environ.get(
        RESUME_BUCKET_ENV
    )


def _embedding_cache_prefix() -> str:
    return os.environ.get(EMBEDDING_CACHE_PREFIX_ENV, DEFAULT_EMBEDDING_CACHE_PREFIX)


def _load_embedding_provider() -> EmbeddingProvider:
    global _embedding_provider

    if _embedding_provider is None:
        provider_name = _semantic_embedding_provider_name()
        if provider_name == BEDROCK_EMBEDDING_PROVIDER:
            _embedding_provider = BedrockEmbeddingProvider(
                _bedrock_runtime_client(),
                _bedrock_embedding_model_id(),
                _bedrock_embedding_dimensions(),
            )
        elif provider_name == LOCAL_EMBEDDING_PROVIDER:
            _embedding_provider = _load_local_embedding_provider()
        else:
            raise ValueError(
                f"{SEMANTIC_EMBEDDING_PROVIDER_ENV} must be "
                f"{BEDROCK_EMBEDDING_PROVIDER!r} or {LOCAL_EMBEDDING_PROVIDER!r}"
            )

    return _embedding_provider


def _load_local_embedding_provider() -> LocalSentenceTransformerEmbeddingProvider:
    model_id = _semantic_model_name()
    try:
        from sentence_transformers import SentenceTransformer
    except ImportError as exc:
        raise RuntimeError(
            "Local semantic matching requires the optional sentence-transformers "
            "package for validation."
        ) from exc

    return LocalSentenceTransformerEmbeddingProvider(SentenceTransformer(model_id), model_id)


def _bedrock_runtime_client() -> Any:
    global bedrock_runtime_client

    if bedrock_runtime_client is None:
        bedrock_runtime_client = boto3.client("bedrock-runtime")

    return bedrock_runtime_client


def _load_embedding_cache() -> S3EmbeddingCache:
    bucket = _embedding_cache_bucket()
    if not bucket:
        raise ValueError(
            f"{EMBEDDING_CACHE_BUCKET_ENV} or {RESUME_BUCKET_ENV} must be configured"
        )

    return S3EmbeddingCache(s3_client, bucket, _embedding_cache_prefix())


def _is_s3_cache_miss(exc: Exception) -> bool:
    error_code = (
        getattr(exc, "response", {})
        .get("Error", {})
        .get("Code")
    )
    return error_code in {"NoSuchKey", "NotFound", "404"} or isinstance(
        exc, FileNotFoundError
    )


def _load_embedding_model() -> Any:
    global _local_embedding_model

    if _local_embedding_model is None:
        try:
            from sentence_transformers import SentenceTransformer
        except ImportError as exc:
            raise RuntimeError(
                "Semantic matching requires the optional sentence-transformers "
                "package for local validation."
            ) from exc

        _local_embedding_model = SentenceTransformer(_semantic_model_name())

    return _local_embedding_model


def _job_description_from_payload(payload: dict[str, Any]) -> str:
    input_fields = [
        name
        for name in ("job_description", "job_description_file", "job_description_url")
        if payload.get(name) is not None
    ]
    if len(input_fields) != 1:
        raise ValueError(
            "provide exactly one of job_description, job_description_file, "
            "or job_description_url"
        )

    if "job_description" in input_fields:
        return payload["job_description"]

    if "job_description_file" in input_fields:
        return _job_description_from_file_payload(payload["job_description_file"])

    return _job_description_from_url(payload["job_description_url"])


def _resume_object_from_payload(payload: dict[str, Any]) -> dict[str, Any] | None:
    if payload.get("resume_text") is None:
        return None

    resume_text = payload["resume_text"]
    if not isinstance(resume_text, str) or not resume_text.strip():
        raise ValueError("resume_text must be a non-empty string when provided")

    return {
        "text": resume_text,
        "source": None,
    }


def _job_description_from_file_payload(file_payload: dict[str, Any]) -> str:
    if not isinstance(file_payload, dict):
        raise ValueError("job_description_file must be an object")

    filename = file_payload.get("filename")
    content = file_payload.get("content")
    if not isinstance(filename, str) or not filename.strip():
        raise ValueError("job_description_file.filename must be a non-empty string")
    if not isinstance(content, str):
        raise ValueError("job_description_file.content must be a string")

    if file_payload.get("is_base64_encoded"):
        body = base64.b64decode(content)
    else:
        body = content.encode("utf-8")

    return _extract_text_from_supported_document(
        body,
        filename,
        SUPPORTED_JOB_FILE_EXTENSIONS,
        "job description",
    )


def _job_description_from_url(url: str) -> str:
    if not isinstance(url, str) or not url.strip():
        raise ValueError("job_description_url must be a non-empty string")

    parsed_url = urllib.parse.urlparse(url)
    _validate_job_description_url(parsed_url)
    request = urllib.request.Request(
        url,
        headers={"User-Agent": "aws-resume-matcher/2.1"},
    )
    with urllib.request.urlopen(
        request,
        timeout=JOB_DESCRIPTION_URL_TIMEOUT_SECONDS,
    ) as response:
        body = response.read(MAX_JOB_DESCRIPTION_URL_BYTES + 1)
        if len(body) > MAX_JOB_DESCRIPTION_URL_BYTES:
            raise ValueError("job_description_url response is too large")

        content_type = response.headers.get("Content-Type", "")

    if "text/html" in content_type.casefold() or _document_extension(url) in {
        ".htm",
        ".html",
    }:
        return _extract_text_from_html(_decode_text(body))

    return _decode_text(body)


def _validate_job_description_url(parsed_url: urllib.parse.ParseResult) -> None:
    if parsed_url.scheme not in {"http", "https"}:
        raise ValueError("job_description_url must use http or https")
    if not parsed_url.hostname:
        raise ValueError("job_description_url must include a hostname")
    if parsed_url.username or parsed_url.password:
        raise ValueError("job_description_url must not include credentials")

    hostname = parsed_url.hostname.casefold()
    if hostname in {"localhost", "127.0.0.1", "::1"}:
        raise ValueError("job_description_url must not point to localhost")

    try:
        ip_address = ipaddress.ip_address(hostname)
    except ValueError:
        return

    if ip_address.is_private or ip_address.is_loopback or ip_address.is_link_local:
        raise ValueError("job_description_url must not point to a private address")


def _extract_text_from_supported_document(
    body: bytes,
    source_name: str,
    supported_extensions: set[str],
    label: str,
) -> str:
    extension = _document_extension(source_name)
    if extension not in supported_extensions:
        supported = ", ".join(sorted(supported_extensions))
        raise ValueError(f"Unsupported {label} file type {extension!r}; use {supported}")

    if extension in {".txt", ".md"}:
        return _decode_text(body)
    if extension == ".pdf":
        return _extract_text_from_pdf(body)
    if extension == ".docx":
        return _extract_text_from_docx(body)

    raise ValueError(f"Unsupported {label} file type {extension!r}")


def _document_extension(source_name: str) -> str:
    path = urllib.parse.urlparse(source_name).path
    _, extension = os.path.splitext(path)
    return extension.casefold()


def _decode_text(body: bytes) -> str:
    return body.decode("utf-8-sig")


def _extract_text_from_pdf(body: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF resume intake requires the pypdf package.") from exc

    reader = PdfReader(BytesIO(body))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_text_from_docx(body: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX resume intake requires the python-docx package.") from exc

    document = Document(BytesIO(body))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)

    return "\n".join(part for part in parts if part)


def _extract_text_from_html(html_text: str) -> str:
    parser = HTMLTextExtractor()
    parser.feed(html_text)
    parser.close()
    return parser.text()


class HTMLTextExtractor(html.parser.HTMLParser):
    def __init__(self):
        super().__init__()
        self.parts = []
        self._ignored_tag_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"script", "style", "noscript"}:
            self._ignored_tag_depth += 1
        if tag in {"br", "p", "div", "li", "section", "article", "h1", "h2", "h3"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"} and self._ignored_tag_depth:
            self._ignored_tag_depth -= 1
        if tag in {"p", "div", "li", "section", "article", "h1", "h2", "h3"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._ignored_tag_depth:
            return
        stripped_data = data.strip()
        if stripped_data:
            self.parts.append(stripped_data)

    def text(self) -> str:
        return re.sub(r"\s+", " ", " ".join(self.parts)).strip()


def _extract_keywords(text: str) -> set[str]:
    normalized_text = _normalize_text(text)
    return {
        keyword
        for raw_keyword in KEYWORD_TOKEN_PATTERN.findall(normalized_text)
        if (keyword := _clean_keyword(raw_keyword))
    }


def _clean_keyword(raw_keyword: str) -> str:
    keyword = raw_keyword.strip(KEYWORD_TRAILING_PUNCTUATION)
    keyword = KEYWORD_ALIASES.get(keyword, keyword)

    if (
        len(keyword) <= 1
        or keyword in STOP_WORDS
        or keyword in CONTRACTION_FRAGMENTS
        or keyword in LOW_VALUE_KEYWORDS
        or keyword[0].isdigit()
        or _is_tokenization_artifact(keyword)
    ):
        return ""

    return keyword


def _is_tokenization_artifact(keyword: str) -> bool:
    return not any(character.isalpha() for character in keyword)


def _normalize_text(text: str) -> str:
    return text.casefold().replace("'", " ").replace("’", " ")


def _response(status_code: int, body: dict[str, Any]) -> dict[str, Any]:
    return {
        "statusCode": status_code,
        "headers": {
            "Content-Type": "application/json",
            "Access-Control-Allow-Origin": "*",
        },
        "body": json.dumps(body),
    }
